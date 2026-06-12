"""
Ingest real failure-state manuals into ChromaDB.

RUN THIS ON THE RUNPOD SERVER (it needs the embedding model, which lives
server-side). It writes through the running Chroma HTTP server.

What it does, fully automatically
---------------------------------
1. Reads every supported file you give it (a folder or individual paths):
       .pdf  .docx  .txt  .md  .csv  .xlsx  .xls  .png  .jpg  .jpeg
   - PDFs: extracts text; if a page is scanned (no text), it OCRs the page
     image with the vision model.
   - Images: OCR'd with the vision model.
   - Spreadsheets: each row becomes a passage.
2. Splits long text into overlapping chunks (big files are fine).
3. Finds the instrument tags (PT-101, LT-104, …) mentioned in each chunk.
4. Stores each chunk once per tag it mentions, so any tag the blueprint
   reader sees can retrieve every manual passage about it.

Usage
-----
    # ingest a whole folder (default: /root/input_docs)
    python3 ingest.py

    # ingest specific files, and wipe the old mock data first
    python3 ingest.py --reset /root/manuals/PT-101.pdf /root/manuals/book.docx
"""
from __future__ import annotations

import argparse
import io
import os
import re
import sys
from pathlib import Path

import config

# NOTE: chromadb is imported lazily inside ingest_files() — importing it at
# module load would pull a heavy dependency into anything that imports this
# module (e.g. the Streamlit app), and on some clients it's slow to import.

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls",
             ".png", ".jpg", ".jpeg"}

CHUNK_CHARS = 1200
CHUNK_OVERLAP = 200
_TAG_RE = re.compile(config.TAG_PATTERN)


# --------------------------------------------------------------------------- #
# Text extraction (one function per format, all return plain text)
# --------------------------------------------------------------------------- #
def _ocr_image_bytes(data: bytes) -> str:
    """OCR an image with the vision model (used for images & scanned PDFs)."""
    import ollama
    client = ollama.Client(host=config.OLLAMA_HOST)
    resp = client.chat(
        model=config.VISION_MODEL,
        messages=[{
            "role": "user",
            "content": (
                "Transcribe ALL text in this document image exactly, including "
                "every alphanumeric tag and code. Preserve line breaks. Output "
                "only the transcribed text."
            ),
            "images": [data],
        }],
    )
    return resp["message"]["content"]


def _extract_pdf(path: Path) -> list[tuple[str, str]]:
    """Per-page passages tagged with their page number (e.g. 'p.12')."""
    import fitz  # PyMuPDF
    out: list[tuple[str, str]] = []
    with fitz.open(path) as doc:
        for pno, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if len(text) < 15:  # likely a scanned page -> OCR it
                pix = page.get_pixmap(dpi=200)
                text = _ocr_image_bytes(pix.tobytes("png"))
            for chunk in chunk_text(text):
                out.append((chunk, f"p.{pno}"))
    return out


def _extract_docx(path: Path) -> list[tuple[str, str]]:
    import docx
    d = docx.Document(str(path))
    lines = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text for c in row.cells))
    # Word has no fixed page numbers; locate by chunk index instead.
    return [(c, f"part {i + 1}") for i, c in enumerate(chunk_text("\n".join(lines)))]


def _spreadsheet_rows(path: Path) -> list[tuple[str, str]]:
    """
    One passage per data row (header prepended), tagged with its row location so
    a tag retrieves exactly its own row, with a precise citation.
    """
    def emit(rows: list[list[str]], sheet: str | None) -> list[tuple[str, str]]:
        rows_idx = [(i, r) for i, r in enumerate(rows, start=1)
                    if any(c.strip() for c in r)]
        if not rows_idx:
            return []
        header = " | ".join(rows_idx[0][1])
        body = rows_idx[1:] if len(rows_idx) > 1 else rows_idx
        out = []
        for rno, r in body:
            loc = f"{sheet}!row {rno}" if sheet else f"row {rno}"
            out.append((f"{header}\n{' | '.join(r)}", loc))
        return out

    if path.suffix.lower() == ".csv":
        import csv
        with open(path, newline="", encoding="utf-8", errors="replace") as fh:
            return emit([list(row) for row in csv.reader(fh)], None)

    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    passages: list[tuple[str, str]] = []
    for ws in wb.worksheets:
        grid = [[("" if c is None else str(c)) for c in row]
                for row in ws.iter_rows(values_only=True)]
        passages += emit(grid, ws.title)
    return passages


def extract_passages(path: Path) -> list[tuple[str, str]]:
    """
    Return a list of (passage_text, location) for a file.

    `location` is a human-readable citation locator: a page number for PDFs,
    a sheet/row for spreadsheets, a part index otherwise (empty if not known).
    Spreadsheets split per-row; everything else is chunked.
    """
    ext = path.suffix.lower()
    if ext in {".csv", ".xlsx", ".xls"}:
        return _spreadsheet_rows(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in {".png", ".jpg", ".jpeg"}:
        return [(c, "") for c in chunk_text(_ocr_image_bytes(path.read_bytes()))]
    if ext in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [(c, "") for c in chunk_text(text)]
    raise ValueError(f"Unsupported file type: {ext}")


# --------------------------------------------------------------------------- #
# Chunking + tag detection
# --------------------------------------------------------------------------- #
def chunk_text(text: str) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= CHUNK_CHARS:
        return [text] if text else []
    chunks, start = [], 0
    while start < len(text):
        end = start + CHUNK_CHARS
        # try to break on a newline near the end for cleaner chunks
        nl = text.rfind("\n", start + CHUNK_CHARS - CHUNK_OVERLAP, end)
        if nl != -1:
            end = nl
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = max(end - CHUNK_OVERLAP, end) if end <= start else end - CHUNK_OVERLAP
        if start <= 0:
            break
    return chunks


def tags_in(text: str) -> list[str]:
    seen: set[str] = set()
    out = []
    for t in _TAG_RE.findall(text.upper()):
        if t not in seen:
            seen.add(t)
            out.append(t)
    return out


# --------------------------------------------------------------------------- #
# Main ingestion
# --------------------------------------------------------------------------- #
def collect_files(paths: list[str], default_dir: str) -> list[Path]:
    targets = paths or [default_dir]
    files: list[Path] = []
    for t in targets:
        p = Path(t)
        if p.is_dir():
            files += [f for f in p.rglob("*") if f.suffix.lower() in SUPPORTED]
        elif p.is_file() and p.suffix.lower() in SUPPORTED:
            files.append(p)
        else:
            print(f"  ! skipping (not a supported file/dir): {t}")
    return sorted(set(files))


def ingest_files(files: list[Path], reset: bool = False,
                 log=print) -> dict:
    """
    Ingest a list of files into ChromaDB. Reusable by both the CLI and the
    Streamlit "Manuals" tab.

    `log` is a callable for progress messages (defaults to print; the UI passes
    its own so messages show in the browser).

    Returns a summary dict:
        {"files", "passages", "entries", "tags": [...], "untagged": [...],
         "total_in_db"}
    """
    import chromadb  # lazy (see note at top of module)
    client = chromadb.HttpClient(host=config.CHROMA_HOST, port=config.CHROMA_PORT)
    if reset:
        try:
            client.delete_collection(config.COLLECTION_NAME)
            log(f"Reset: cleared collection '{config.COLLECTION_NAME}'.")
        except Exception:  # noqa: BLE001
            pass
    collection = client.get_or_create_collection(config.COLLECTION_NAME)

    total_passages = total_entries = 0
    all_tags: set[str] = set()
    untagged: list[str] = []

    for f in files:
        log(f"• {f.name}")
        try:
            passages = extract_passages(f)
        except Exception as e:  # noqa: BLE001
            log(f"    ✗ failed to read: {e}")
            continue

        docs, metas, ids = [], [], []
        file_tag_hits = 0
        for ci, (passage, loc) in enumerate(passages):
            for tag in tags_in(passage):
                docs.append(passage)
                metas.append({"tag": tag, "source": f.name, "chunk": ci,
                              "loc": loc})
                ids.append(f"{f.name}::{ci}::{tag}")
                all_tags.add(tag)
                file_tag_hits += 1

        if docs:
            import embeddings
            collection.upsert(documents=docs, embeddings=embeddings.embed_texts(docs),
                              metadatas=metas, ids=ids)
        total_passages += len(passages)
        total_entries += len(docs)
        if file_tag_hits == 0:
            untagged.append(f.name)
        log(f"    {len(passages)} passage(s), {file_tag_hits} tag-entr(ies)")

    return {
        "files": len(files),
        "passages": total_passages,
        "entries": total_entries,
        "tags": sorted(all_tags),
        "untagged": untagged,
        "total_in_db": collection.count(),
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="Ingest manuals into ChromaDB.")
    ap.add_argument("paths", nargs="*", help="Files or folders to ingest.")
    ap.add_argument("--dir", default=os.getenv("INPUT_DOCS_DIR", "./input_docs"),
                    help="Default folder to ingest if no paths given "
                         "(env: INPUT_DOCS_DIR).")
    ap.add_argument("--reset", action="store_true",
                    help="Delete existing collection contents first.")
    args = ap.parse_args()

    files = collect_files(args.paths, args.dir)
    if not files:
        print(f"No supported files found. Put manuals in {args.dir} or pass paths.")
        print(f"Supported: {', '.join(sorted(SUPPORTED))}")
        return 1

    print(f"Found {len(files)} file(s) to ingest.\n")
    s = ingest_files(files, reset=args.reset)

    print("\n── Summary ─────────────────────────────")
    print(f"Files ingested : {s['files']}")
    print(f"Passages total : {s['passages']}")
    print(f"Tag entries    : {s['entries']}")
    print(f"Unique tags    : {len(s['tags'])}")
    if s["tags"]:
        print(f"Tags           : {', '.join(s['tags'])}")
    if s["untagged"]:
        print(f"\n⚠ No instrument tags found in: {', '.join(s['untagged'])}")
        print("  (Those files were read but won't be retrievable by tag.)")
    print(f"\nCollection now holds {s['total_in_db']} entries.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
